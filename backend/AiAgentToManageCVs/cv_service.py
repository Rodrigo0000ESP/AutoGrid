import os
import io
from typing import Dict, Any
import pdfplumber
from docx import Document
from openai import OpenAI


class CVService:
    def __init__(self):
        self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
    def generate_cv(self, cv_file_path: str, job_description: str, job_skills: str) -> Dict[str, Any]:
        """Generate CV based on file type"""
        if cv_file_path.endswith(".pdf"):
            return self._generate_pdf_cv(cv_file_path, job_description, job_skills)
        elif cv_file_path.endswith(".docx"):
            return self._generate_docx_cv(cv_file_path, job_description, job_skills)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")

    def _generate_pdf_cv(self, cv_file_path: str, job_description: str, job_skills: str) -> Dict[str, Any]:
        """For PDF: Extract text and return AI suggestions"""
        # Extract text from PDF
        cv_text = ""
        with pdfplumber.open(cv_file_path) as pdf:
            for page in pdf.pages:
                cv_text += page.extract_text() + "\n"
        
        # Generate suggestions using AI
        suggestion = self.generate_suggestion(cv_text, job_description, job_skills)
        
        return {
            "type": "suggestion",
            "content": suggestion,
            "original_text": cv_text
        }

    def _generate_docx_cv(self, cv_file_path: str, job_description: str, job_skills: str) -> Dict[str, Any]:
        """For DOCX: Generate new CV maintaining original style"""
        # Extract text from DOCX
        doc = Document(cv_file_path)
        cv_text = "\n".join([para.text for para in doc.paragraphs])
        
        # Generate new CV content using AI
        new_content = self.generate_new_cv(cv_text, job_description, job_skills)
        
        # Create new DOCX with updated content but same style
        new_doc = Document(cv_file_path)  # Copy original document with styles
        
        # Replace content while maintaining formatting
        content_lines = new_content.split("\n")
        for i, paragraph in enumerate(new_doc.paragraphs):
            if i < len(content_lines):
                paragraph.text = content_lines[i]
        
        # Save to BytesIO
        output = io.BytesIO()
        new_doc.save(output)
        output.seek(0)
        
        return {
            "type": "document",
            "content": output,
            "filename": "generated_cv.docx"
        }

    def generate_suggestion(self, cv_text: str, job_description: str, job_skills: str) -> str:
        """Generate AI suggestions for CV improvement"""
        prompt = f"""You are an expert CV consultant. Analyze the following CV and provide specific suggestions to tailor it for this job opportunity.

CV Content:
{cv_text}

Job Description:
{job_description}

Required Skills:
{job_skills}

Provide a detailed analysis with:
1. Key strengths that match the job
2. Skills to emphasize more
3. Experience to highlight
4. Specific improvements to make
5. Keywords to add

Be specific and actionable."""

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an expert CV consultant who helps people optimize their CVs for specific job opportunities."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=2000
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error generating suggestions: {str(e)}"

    def generate_new_cv(self, cv_text: str, job_description: str, job_skills: str) -> str:
        """Generate new CV content tailored to job offer"""
        prompt = f"""You are an expert CV writer. Rewrite the following CV to perfectly match this job opportunity while maintaining the person's authentic experience and achievements.

Original CV:
{cv_text}

Job Description:
{job_description}

Required Skills:
{job_skills}

Rewrite the CV to:
1. Emphasize relevant experience and skills
2. Use keywords from the job description
3. Highlight achievements that match job requirements
4. Maintain the same structure and sections
5. Keep all information truthful and accurate

Provide the complete rewritten CV text, maintaining professional formatting."""

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an expert CV writer who tailors CVs to specific job opportunities while maintaining authenticity."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=3000
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error generating CV: {str(e)}"
